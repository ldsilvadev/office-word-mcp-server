"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import find_and_replace_text, insert_header_near_text, insert_numbered_list_near_text, insert_line_or_paragraph_near_text, replace_paragraph_block_below_header, replace_block_between_manual_anchors
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def add_heading(filename: str, text: str, level: int = 1,
                      font_name: Optional[str] = None, font_size: Optional[int] = None,
                      bold: Optional[bool] = None, italic: Optional[bool] = None,
                      border_bottom: bool = False) -> str:
    """Add a heading to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
        font_name: Font family (e.g., 'Helvetica')
        font_size: Font size in points (e.g., 14)
        bold: True/False for bold text
        italic: True/False for italic text
        border_bottom: True to add bottom border (for section headers)
    """
    filename = ensure_docx_extension(filename)

    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"

    # Validate level range
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        # Ensure heading styles exist
        ensure_heading_style(doc)

        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            heading = doc.add_paragraph(text)
            heading.style = doc.styles['Normal']
            if heading.runs:
                run = heading.runs[0]
                run.bold = True
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)

        # Apply formatting to all runs in the heading
        if any([font_name, font_size, bold is not None, italic is not None]):
            for run in heading.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic

        # Add bottom border if requested
        if border_bottom:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            pPr = heading._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')  # 0.5pt border
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')

            pBdr.append(bottom)
            pPr.append(pBdr)

        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


async def add_paragraph(filename: str, text: str, style: Optional[str] = None,
                        font_name: Optional[str] = None, font_size: Optional[int] = None,
                        bold: Optional[bool] = None, italic: Optional[bool] = None,
                        color: Optional[str] = None) -> str:
    """Add a paragraph to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
        font_name: Font family (e.g., 'Helvetica', 'Times New Roman')
        font_size: Font size in points (e.g., 14, 36)
        bold: True/False for bold text
        italic: True/False for italic text
        color: RGB color as hex string (e.g., '000000' for black)
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        
        # If no style provided, try to inherit from the last paragraph
        if style is None and len(doc.paragraphs) > 0:
            last_para = doc.paragraphs[-1]
            if last_para.style:
                # Don't inherit Heading styles automatically, as we usually want body text after a heading
                if not last_para.style.name.startswith('Heading'):
                    style = last_para.style.name
                    
        paragraph = doc.add_paragraph(text)

        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"

        # Apply formatting to all runs in the paragraph
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    # Remove any '#' prefix if present
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)

        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def add_table_of_contents(filename: str, title: str = "Table of Contents", max_level: int = 3) -> str:
    """Add a table of contents to a Word document based on heading styles.
    
    Args:
        filename: Path to the Word document
        title: Optional title for the table of contents
        max_level: Maximum heading level to include (1-9)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Ensure max_level is within valid range
        max_level = max(1, min(max_level, 9))
        
        doc = Document(filename)
        
        # Collect headings and their positions
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph style is a heading
            if paragraph.style and paragraph.style.name.startswith('Heading '):
                try:
                    # Extract heading level from style name
                    level = int(paragraph.style.name.split(' ')[1])
                    if level <= max_level:
                        headings.append({
                            'level': level,
                            'text': paragraph.text,
                            'position': i
                        })
                except (ValueError, IndexError):
                    # Skip if heading level can't be determined
                    pass
        
        if not headings:
            return f"No headings found in document {filename}. Table of contents not created."
        
        # Create a new document with the TOC
        toc_doc = Document()
        
        # Add title
        if title:
            toc_doc.add_heading(title, level=1)
        
        # Add TOC entries
        for heading in headings:
            # Indent based on level (using tab characters)
            indent = '    ' * (heading['level'] - 1)
            toc_doc.add_paragraph(f"{indent}{heading['text']}")
        
        # Add page break
        toc_doc.add_page_break()
        
        # Get content from original document
        for paragraph in doc.paragraphs:
            p = toc_doc.add_paragraph(paragraph.text)
            # Copy style if possible
            try:
                if paragraph.style:
                    p.style = paragraph.style.name
            except:
                pass
        
        # Copy tables
        for table in doc.tables:
            # Create a new table with the same dimensions
            new_table = toc_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_table.cell(i, j).text = paragraph.text
        
        # Save the new document with TOC
        toc_doc.save(filename)
        
        return f"Table of contents with {len(headings)} entries added to {filename}"
    except Exception as e:
        return f"Failed to add table of contents: {str(e)}"


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

async def insert_header_near_text_tool(filename: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_header_near_text(filename, target_text, header_title, position, header_style, target_paragraph_index)

async def insert_numbered_list_near_text_tool(filename: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """Insert a bulleted or numbered list before or after the target paragraph. Specify by text or paragraph index."""
    return insert_numbered_list_near_text(filename, target_text, list_items, position, target_paragraph_index, bullet_type)

async def insert_line_or_paragraph_near_text_tool(filename: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """Insert a new line or paragraph (with specified or matched style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_line_or_paragraph_near_text(filename, target_text, line_text, position, line_style, target_paragraph_index)

async def replace_paragraph_block_below_header_tool(filename: str, header_text: str, new_paragraphs: list, detect_block_end_fn=None) -> str:
    """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
    return replace_paragraph_block_below_header(filename, header_text, new_paragraphs, detect_block_end_fn)

async def replace_block_between_manual_anchors_tool(filename: str, start_anchor_text: str, new_paragraphs: list, end_anchor_text: str = None, match_fn=None, new_paragraph_style: str = None) -> str:
    """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
    return replace_block_between_manual_anchors(filename, start_anchor_text, new_paragraphs, end_anchor_text, match_fn, new_paragraph_style)


async def edit_paragraph_text(filename: str, paragraph_index: int, new_text: str) -> str:
    """Edit the text of a specific paragraph in a Word document.

    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to edit (0-based)
        new_text: New text for the paragraph
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)

        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."

        # Update the paragraph text
        paragraph = doc.paragraphs[paragraph_index]
        paragraph.text = new_text

        doc.save(filename)
        return f"Paragraph at index {paragraph_index} updated successfully."
    except Exception as e:
        return f"Failed to edit paragraph: {str(e)}"


async def insert_text_inline(filename: str, search_text: str, text_to_insert: str, position: str = 'after') -> str:
    """Insert text inline (same paragraph) before or after a specific text.
    
    Args:
        filename: Path to the Word document
        search_text: Text to search for
        text_to_insert: Text to insert
        position: 'before' or 'after' the search text
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
        
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
        
    try:
        doc = Document(filename)
        modified_count = 0
        
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                # We found the text. Now we need to modify the runs to preserve formatting if possible,
                # or just modify the text directly if simple.
                # For simplicity and robustness in this version, we'll do a text replacement on the paragraph.
                # This might reset some run-level formatting in the specific paragraph, but ensures the text is inline.
                
                original_text = paragraph.text
                if position == 'after':
                    new_text = original_text.replace(search_text, search_text + text_to_insert)
                else: # before
                    new_text = original_text.replace(search_text, text_to_insert + search_text)
                
                if new_text != original_text:
                    paragraph.text = new_text
                    modified_count += 1
        
        if modified_count > 0:
            doc.save(filename)
            return f"Inserted text '{text_to_insert}' {position} '{search_text}' in {modified_count} location(s)."
        else:
            return f"Text '{search_text}' not found in document."
            
    except Exception as e:
        return f"Failed to insert text inline: {str(e)}"


def _normalize_table_data(table_data: List[Any]) -> List[Dict[str, Any]]:
    """Convert table_data to list of dictionaries format.
    
    Accepts:
        - List of dictionaries: [{"Col1": "val1"}, ...] -> returns as-is
        - 2D list: [["Header1", "Header2"], ["val1", "val2"], ...] -> converts to dict format
    
    Returns:
        List of dictionaries with headers as keys
    """
    if not table_data or len(table_data) == 0:
        return []
    
    # Already in dict format
    if isinstance(table_data[0], dict):
        return table_data
    
    # Convert 2D list to dict format
    # First row is headers, rest are data
    if len(table_data) < 2:
        # Only headers, no data - return empty with headers as a single row
        headers = [str(h) for h in table_data[0]]
        return [{h: "" for h in headers}]
    
    headers = [str(h) for h in table_data[0]]
    result = []
    for row in table_data[1:]:
        row_dict = {}
        for j, header in enumerate(headers):
            row_dict[header] = str(row[j]) if j < len(row) else ""
        result.append(row_dict)
    
    return result


def _create_table_in_doc(doc: Document, table_data: List[Dict[str, Any]]) -> Any:
    """Create a table in the document from list of dictionaries.
    
    Args:
        doc: Document object
        table_data: List of dictionaries where keys become headers
        
    Returns:
        The created table object
    """
    if not table_data or len(table_data) == 0:
        return None
    
    headers = list(table_data[0].keys())
    num_cols = len(headers)
    num_rows = len(table_data) + 1  # +1 for header row
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    
    # Fill header row with bold text
    header_row = table.rows[0]
    for j, header in enumerate(headers):
        cell = header_row.cells[j]
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Fill data rows
    for i, row_data in enumerate(table_data):
        row = table.rows[i + 1]
        for j, header in enumerate(headers):
            cell = row.cells[j]
            cell.text = str(row_data.get(header, ""))
    
    return table


async def add_section_with_inherited_formatting(filename: str, title: str, paragraph_text: Optional[str] = None, 
                                              table_data: Optional[List[Any]] = None) -> str:
    """Add a new section (Title + optional Content) inheriting the style of the last heading.
    
    Args:
        filename: Path to the Word document
        title: Text for the new section title
        paragraph_text: Optional text for a paragraph below the title
        table_data: Table data in one of these formats:
                   - List of dictionaries: [{"Col1": "val1", "Col2": "val2"}, ...] (keys become headers)
                   - 2D list: [["Header1", "Header2"], ["val1", "val2"], ...] (first row is header)
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
        
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
        
    try:
        doc = Document(filename)
        
        # Find the last heading style and body style
        target_style = "Heading 1"
        body_style = None
        
        for i in range(len(doc.paragraphs) - 1, -1, -1):
            para = doc.paragraphs[i]
            if para.style and para.style.name.startswith("Heading"):
                target_style = para.style.name
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    if next_para.style and not next_para.style.name.startswith("Heading"):
                        body_style = next_para.style.name
                break
        
        # Add the title with inherited heading style
        title_para = doc.add_paragraph(title, style=target_style)
        last_element = title_para._p
        
        # Add optional paragraph with proper formatting
        if paragraph_text:
            para = doc.add_paragraph()
            para.text = paragraph_text
            if body_style:
                para.style = body_style
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Cm(1.25)
            last_element = para._p
            
        # Add optional table using shared helper
        if table_data and len(table_data) > 0:
            normalized_data = _normalize_table_data(table_data)
            if normalized_data:
                _create_table_in_doc(doc, normalized_data)
                            
        doc.save(filename)
        
        details = []
        if paragraph_text: details.append(f"paragraph (style: {body_style or 'default'})")
        if table_data: details.append("table")
        content_str = f" with {', '.join(details)}" if details else ""
        
        return f"Added section '{title}' (inherited style '{target_style}'){content_str} to {filename}"
            
    except Exception as e:
        return f"Failed to add section: {str(e)}"


# ============================================================================
# SECTION EDITING TOOLS - Edit sections by number without regenerating document
# ============================================================================

async def list_document_sections(filename: str) -> str:
    """List all sections (headings) in a Word document with their numbers and content preview.
    
    Args:
        filename: Path to the Word document
        
    Returns:
        JSON string with list of sections: [{number, title, paragraph_preview, paragraph_index}]
    """
    import json
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        sections = []
        current_section = None
        section_number = 0
        
        for i, para in enumerate(doc.paragraphs):
            # Check if this is a heading (section title)
            if para.style and para.style.name.startswith("Heading"):
                # Save previous section if exists
                if current_section:
                    sections.append(current_section)
                
                section_number += 1
                current_section = {
                    "number": section_number,
                    "title": para.text.strip(),
                    "title_index": i,
                    "heading_style": para.style.name,
                    "paragraphs": [],
                    "paragraph_indices": []
                }
            elif current_section and para.text.strip():
                # This is content under the current section
                current_section["paragraphs"].append(para.text.strip()[:100] + "..." if len(para.text) > 100 else para.text.strip())
                current_section["paragraph_indices"].append(i)
        
        # Don't forget the last section
        if current_section:
            sections.append(current_section)
        
        result = {
            "total_sections": len(sections),
            "sections": sections
        }
        
        return json.dumps(result, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return f"Failed to list sections: {str(e)}"


async def get_section_content(filename: str, section_number: int) -> str:
    """Get the full content of a specific section by its number.
    
    Args:
        filename: Path to the Word document
        section_number: Section number (1-based, e.g., 1 for first section)
        
    Returns:
        JSON with section title and full paragraph content
    """
    import json
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        current_section_num = 0
        section_start_idx = None
        section_title = None
        section_paragraphs = []
        section_para_indices = []
        
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name.startswith("Heading"):
                # If we were collecting a section and found a new heading, stop
                if current_section_num == section_number:
                    break
                    
                current_section_num += 1
                if current_section_num == section_number:
                    section_start_idx = i
                    section_title = para.text.strip()
            elif current_section_num == section_number and para.text.strip():
                section_paragraphs.append(para.text)
                section_para_indices.append(i)
        
        if section_title is None:
            return f"Section {section_number} not found. Document has {current_section_num} sections."
        
        result = {
            "section_number": section_number,
            "title": section_title,
            "title_index": section_start_idx,
            "content": "\n\n".join(section_paragraphs),
            "paragraph_indices": section_para_indices,
            "paragraph_count": len(section_paragraphs)
        }
        
        return json.dumps(result, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return f"Failed to get section content: {str(e)}"


async def append_to_section(filename: str, section_number: int, text_to_append: str) -> str:
    """Append text to the end of a specific section (before the next section starts).
    
    This is the PREFERRED way to add content to a section without regenerating the document.
    
    Args:
        filename: Path to the Word document
        section_number: Section number (1-based, e.g., 4 for fourth section)
        text_to_append: Text to append to the section
        
    Returns:
        Success message or error
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        doc = Document(filename)
        current_section_num = 0
        last_para_in_section_idx = None
        section_title = None
        
        # Find the last paragraph of the target section
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name.startswith("Heading"):
                current_section_num += 1
                if current_section_num == section_number:
                    section_title = para.text.strip()
                    last_para_in_section_idx = i  # Start with the heading itself
                elif current_section_num > section_number:
                    # We've reached the next section, stop
                    break
            elif current_section_num == section_number:
                # Update last paragraph index for this section
                last_para_in_section_idx = i
        
        if section_title is None:
            return f"Section {section_number} not found. Document has {current_section_num} sections."
        
        if last_para_in_section_idx is None:
            return f"Could not find insertion point for section {section_number}"
        
        # Get the paragraph to insert after
        target_para = doc.paragraphs[last_para_in_section_idx]
        
        # Create new paragraph element
        new_para = doc.add_paragraph()
        new_para.text = text_to_append
        
        # Try to match the style of existing paragraphs in the section
        # Look for a non-heading paragraph style in the section
        body_style = None
        for i in range(last_para_in_section_idx, -1, -1):
            para = doc.paragraphs[i]
            if para.style and not para.style.name.startswith("Heading"):
                body_style = para.style
                break
        
        if body_style:
            new_para.style = body_style
        
        # Set paragraph formatting
        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_para.paragraph_format.space_after = Pt(6)
        new_para.paragraph_format.first_line_indent = Cm(1.25)
        
        # Move the new paragraph to the correct position (after target_para)
        target_para._p.addnext(new_para._p)
        
        doc.save(filename)
        
        return f"✅ Text appended to section {section_number} ('{section_title}'). The paragraph was inserted after index {last_para_in_section_idx}."
        
    except Exception as e:
        return f"Failed to append to section: {str(e)}"


async def replace_section_content(filename: str, section_number: int, new_content: str, keep_title: bool = True) -> str:
    """Replace all content in a specific section with new content.
    
    Args:
        filename: Path to the Word document
        section_number: Section number (1-based)
        new_content: New content for the section (replaces all paragraphs)
        keep_title: If True, keeps the section title unchanged (default: True)
        
    Returns:
        Success message or error
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        doc = Document(filename)
        current_section_num = 0
        section_title = None
        section_title_idx = None
        paragraphs_to_remove = []
        
        # Find all paragraphs in the target section
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name.startswith("Heading"):
                current_section_num += 1
                if current_section_num == section_number:
                    section_title = para.text.strip()
                    section_title_idx = i
                elif current_section_num > section_number:
                    break
            elif current_section_num == section_number:
                paragraphs_to_remove.append(i)
        
        if section_title is None:
            return f"Section {section_number} not found."
        
        # Remove existing paragraphs (in reverse order to maintain indices)
        for idx in reversed(paragraphs_to_remove):
            para = doc.paragraphs[idx]
            p = para._element
            p.getparent().remove(p)
        
        # Get the title paragraph to insert after
        title_para = doc.paragraphs[section_title_idx]
        
        # Add new content paragraphs
        new_paragraphs = new_content.split("\n\n")
        last_element = title_para._p
        
        for para_text in new_paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
                
            new_para = doc.add_paragraph()
            new_para.text = para_text
            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            new_para.paragraph_format.space_after = Pt(6)
            new_para.paragraph_format.first_line_indent = Cm(1.25)
            
            # Move to correct position
            new_para._p.getparent().remove(new_para._p)
            last_element.addnext(new_para._p)
            last_element = new_para._p
        
        doc.save(filename)
        
        return f"✅ Section {section_number} ('{section_title}') content replaced successfully."
        
    except Exception as e:
        return f"Failed to replace section content: {str(e)}"


async def edit_section_title(filename: str, section_number: int, new_title: str) -> str:
    """Edit the title of a specific section.
    
    Args:
        filename: Path to the Word document
        section_number: Section number (1-based)
        new_title: New title for the section
        
    Returns:
        Success message or error
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        doc = Document(filename)
        current_section_num = 0
        
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                current_section_num += 1
                if current_section_num == section_number:
                    old_title = para.text
                    para.text = new_title
                    doc.save(filename)
                    return f"✅ Section {section_number} title changed from '{old_title}' to '{new_title}'."
        
        return f"Section {section_number} not found. Document has {current_section_num} sections."
        
    except Exception as e:
        return f"Failed to edit section title: {str(e)}"


async def append_table_to_section(filename: str, section_number: int, table_data: List[Any], 
                                   paragraph_before: str = None, paragraph_after: str = None) -> str:
    """Append a table to the end of a specific section.
    
    This tool inserts a table at the end of a section, optionally with text before and/or after.
    
    Args:
        filename: Path to the Word document
        section_number: Section number (1-based, e.g., 4 for fourth section)
        table_data: Table data in one of these formats:
                   - List of dictionaries (PREFERRED): [{"Cargo": "Diretor", "Limite": "1000"}, ...]
                   - 2D list: [["Cargo", "Limite"], ["Diretor", "1000"], ...]
        paragraph_before: Optional text to add BEFORE the table
        paragraph_after: Optional text to add AFTER the table
        
    Returns:
        Success message or error
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    if not table_data or len(table_data) == 0:
        return "Error: table_data cannot be empty"
    
    # Normalize table data to dict format
    normalized_data = _normalize_table_data(table_data)
    if not normalized_data:
        return "Error: Could not process table_data"
    
    try:
        doc = Document(filename)
        current_section_num = 0
        last_para_in_section_idx = None
        section_title = None
        
        # Find the last paragraph of the target section
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name.startswith("Heading"):
                current_section_num += 1
                if current_section_num == section_number:
                    section_title = para.text.strip()
                    last_para_in_section_idx = i
                elif current_section_num > section_number:
                    break
            elif current_section_num == section_number:
                last_para_in_section_idx = i
        
        if section_title is None:
            return f"Section {section_number} not found. Document has {current_section_num} sections."
        
        if last_para_in_section_idx is None:
            return f"Could not find insertion point for section {section_number}"
        
        # Get the paragraph to insert after
        target_para = doc.paragraphs[last_para_in_section_idx]
        last_element = target_para._p
        
        # Find body style
        body_style = None
        for i in range(last_para_in_section_idx, -1, -1):
            para = doc.paragraphs[i]
            if para.style and not para.style.name.startswith("Heading"):
                body_style = para.style
                break
        
        # Add paragraph before table if provided
        if paragraph_before:
            para_before = doc.add_paragraph()
            para_before.text = paragraph_before
            if body_style:
                para_before.style = body_style
            para_before.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para_before.paragraph_format.space_after = Pt(6)
            para_before.paragraph_format.first_line_indent = Cm(1.25)
            
            para_before._p.getparent().remove(para_before._p)
            last_element.addnext(para_before._p)
            last_element = para_before._p
        
        # Create table using shared helper
        table = _create_table_in_doc(doc, normalized_data)
        
        # Move table to correct position
        tbl = table._tbl
        tbl.getparent().remove(tbl)
        last_element.addnext(tbl)
        last_element = tbl
        
        # Add paragraph after table if provided
        if paragraph_after:
            para_after = doc.add_paragraph()
            para_after.text = paragraph_after
            if body_style:
                para_after.style = body_style
            para_after.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para_after.paragraph_format.space_after = Pt(6)
            para_after.paragraph_format.first_line_indent = Cm(1.25)
            
            para_after._p.getparent().remove(para_after._p)
            last_element.addnext(para_after._p)
        
        doc.save(filename)
        
        num_cols = len(normalized_data[0].keys())
        result_parts = [f"✅ Table ({len(normalized_data)} rows x {num_cols} cols) inserted in section {section_number} ('{section_title}')"]
        if paragraph_before:
            result_parts.append("with text before")
        if paragraph_after:
            result_parts.append("with text after")
        
        return " ".join(result_parts) + "."
        
    except Exception as e:
        return f"Failed to append table to section: {str(e)}"
