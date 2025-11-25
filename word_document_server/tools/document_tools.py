"""
Document creation and manipulation tools for Word Document Server.
"""
import os
import json
import asyncio
from typing import Dict, List, Optional, Any
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docxtpl import DocxTemplate
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, create_document_copy
from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure, get_document_xml, insert_header_near_text, insert_line_or_paragraph_near_text
from word_document_server.core.styles import ensure_heading_style, ensure_table_style

logging.basicConfig(
    filename='mcp_debug.log', 
    level=logging.INFO, 
    format='%(asctime)s - %(levelname)s - %(message)s',
    force=True # Garante que vai sobrescrever configs anteriores
)

async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    filename = ensure_docx_extension(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get information about a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    return extract_document_text(filename)


async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)
    
    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"


async def merge_documents(target_filename: str, source_filenames: List[str], add_page_breaks: bool = True) -> str:
    """Merge multiple Word documents into a single document.
    
    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
    """
    from word_document_server.core.tables import copy_table
    
    target_filename = ensure_docx_extension(target_filename)
    
    # Check if target file is writeable
    is_writeable, error_message = check_file_writeable(target_filename)
    if not is_writeable:
        return f"Cannot create target document: {error_message}"
    
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)
    
    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"
    
    try:
        # Create a new document for the merged result
        target_doc = Document()
        
        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc = Document(doc_filename)
            
            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()
            
            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles['Normal']  # Default style
                
                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass
                
                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size
            
            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)
        
        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"


async def get_document_xml_tool(filename: str) -> str:
    """Get the raw XML structure of a Word document."""
    return get_document_xml(filename)


async def fill_document_template(template_path: str, output_path: str, data_json: str) -> str:
    """Fill a Word document template using docxtpl (Jinja2) with provided data.
    
    This function allows you to populate Word templates with dynamic content including:
    - Simple variables (e.g., {{assunto}}, {{codigo}})
    - Loops for tables (e.g., {% for item in secao %})
    - Conditional content
    - Header/footer variables
    
    Args:
        template_path: Path to the template Word document (.docx)
        output_path: Path where the filled document will be saved
        data_json: JSON string containing the data to fill the template.
                   Must be a valid JSON object that matches template variables.
                   Example: '{"assunto": "Test", "codigo": "123", "secao": [{"nome": "Item 1"}]}'
    
    Returns:
        Success message or error description
    
    Example:
        data = {
            "assunto": "Monthly Report",
            "codigo": "RPT-2024-001",
            "items": [
                {"name": "Item 1", "value": 100},
                {"name": "Item 2", "value": 200}
            ]
        }
        await fill_document_template("template.docx", "output.docx", json.dumps(data))
    """
    template_path = ensure_docx_extension(template_path)
    output_path = ensure_docx_extension(output_path)
    
    # Check if template exists
    if not os.path.exists(template_path):
        return f"Template file {template_path} does not exist"
    
    # Check if output file is writeable
    is_writeable, error_message = check_file_writeable(output_path)
    if not is_writeable:
        return f"Cannot write to output file: {error_message}"
    
    # Parse JSON data
    try:
        context_data = json.loads(data_json)
        if not isinstance(context_data, dict):
            return "Error: data_json must be a JSON object (dictionary), not an array or primitive type"
    except json.JSONDecodeError as e:
        return f"Error parsing JSON data: {str(e)}. Please provide valid JSON string."
    
    # Define the blocking I/O operation
    def _fill_template():
        """Internal function to perform blocking template operations."""
        try:
            # Load the template
            doc = DocxTemplate(template_path)
            
            # Render the template with provided context
            doc.render(context_data)
            
            # Save the filled document
            doc.save(output_path)
            
            return True, f"Template filled successfully. Output saved to {output_path}"
        except Exception as e:
            return False, f"Error filling template: {str(e)}"
    
    # Execute the blocking operation in a thread pool to avoid blocking the async loop
    try:
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            success, message = await loop.run_in_executor(executor, _fill_template)
            return message
    except Exception as e:
        return f"Failed to fill document template: {str(e)}"


def _replace_in_runs(paragraph, placeholder, value):
    """Replace placeholder in paragraph runs, handling split placeholders."""
    full_text = paragraph.text
    
    if placeholder not in full_text:
        return False
    
    # Try simple replacement in runs first
    replaced = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            replaced = True
    
    if replaced:
        return True
    
    # Placeholder is split across runs - rebuild paragraph
    new_text = full_text.replace(placeholder, value)
    
    if not paragraph.runs:
        return False
    
    # Keep first run's formatting
    first_run = paragraph.runs[0]
    bold = first_run.bold
    italic = first_run.italic
    underline = first_run.underline
    font_size = first_run.font.size
    font_name = first_run.font.name
    
    # Clear all runs
    for run in paragraph.runs:
        run.text = ""
    
    # Set new text in first run
    first_run.text = new_text
    
    # Restore formatting
    if bold is not None:
        first_run.bold = bold
    if italic is not None:
        first_run.italic = italic
    if underline is not None:
        first_run.underline = underline
    if font_size:
        first_run.font.size = font_size
    if font_name:
        first_run.font.name = font_name
    
    return True

def _process_dynamic_table_placeholder(doc, context_data, placeholder_key="tabela_dinamica"):
    import logging 
    
    target_para = None
    # Sintaxe correta para achar {{tabela_dinamica}}
    clean_key = f"{{{{{placeholder_key}}}}}" 
    
    # 1. Encontrar o parágrafo alvo (onde a tabela vai ficar)
    for para in doc.paragraphs:
        if clean_key in para.text.replace(" ", ""):
            target_para = para
            break
            
    if not target_para:
        return 

    # 2. Validar dados
    data = context_data.get(placeholder_key)
    if not data or not isinstance(data, list) or len(data) == 0 or not isinstance(data[0], dict):
        # Remove o placeholder se não houver dados
        p = target_para._element
        p.getparent().remove(p)
        return

    # 3. Criar a Tabela
    headers = list(data[0].keys())
    rows = len(data)
    cols = len(headers)
    
    # Cria a tabela (o Word joga ela pro final por padrão)
    table = doc.add_table(rows=rows + 1, cols=cols)
    
    # Aplica estilo (já resolvido por você com try/except se necessário)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    
    # Preencher Cabeçalho e Dados
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header).upper().replace("_", " ")
        if hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
    for i, item in enumerate(data):
        row_cells = table.rows[i + 1].cells
        for j, key in enumerate(headers):
            val = item.get(key, "")
            row_cells[j].text = str(val)
            
    # 4. A MÁGICA DE POSICIONAMENTO:
    # Move a tabela do final do documento para logo após o parágrafo alvo
    target_para._p.addnext(table._tbl)
    
    # 5. Remover o parágrafo do placeholder
    p = target_para._element
    p.getparent().remove(p)

async def fill_document_simple(template_path: str, output_path: str, data_json: str) -> str:
    """Fill a Word document template using simple text replacement (python-docx).
    
    This function provides better formatting preservation than docxtpl by using
    direct text replacement. It maintains ALL original formatting including:
    - Paragraph styles and spacing
    - Font formatting (size, color, bold, italic)
    - Alignment and indentation
    - Line spacing
    
    Supports:
    - Simple variables: {{variavel}}
    - Loop expansion: {{LOOP:secao}} with {{titulo}} and {{paragrafo}} inside
    
    Args:
        template_path: Path to the template Word document (.docx)
        output_path: Path where the filled document will be saved
        data_json: JSON string containing the data to fill the template.
                   Example: '{"assunto": "Test", "secao": [{"titulo": "T1", "paragrafo": "P1"}]}'
    
    Returns:
        Success message or error description
    
    Example:
        data = {
            "assunto": "Política de Home Office",
            "codigo": "FIERGS-001",
            "secao": [
                {"titulo": "1. Objetivo", "paragrafo": "Esta política..."},
                {"titulo": "2. Elegibilidade", "paragrafo": "O regime..."}
            ]
        }
        await fill_document_simple("template.docx", "output.docx", json.dumps(data))
    """
    template_path = ensure_docx_extension(template_path)
    output_path = ensure_docx_extension(output_path)
    
    # Check if template exists
    if not os.path.exists(template_path):
        return f"Template file {template_path} does not exist"
    
    # Check if output file is writeable
    is_writeable, error_message = check_file_writeable(output_path)
    if not is_writeable:
        return f"Cannot write to output file: {error_message}"
    
    # Parse JSON data
    try:
        context_data = json.loads(data_json)
        if not isinstance(context_data, dict):
            return "Error: data_json must be a JSON object (dictionary), not an array or primitive type"
    except json.JSONDecodeError as e:
        return f"Error parsing JSON data: {str(e)}. Please provide valid JSON string."
    

    def _fill_template():
        """Internal function to perform blocking template operations."""
        try:
            from copy import deepcopy
            
            doc = Document(template_path)
            
            # STEP 1: Find and process loop (FLEXÍVEL)
            loop_marker_idx = None
            titulo_idx = None
            paragrafo_idx = None
            loop_var = None
            
            # Detectar marcadores
            for i, para in enumerate(doc.paragraphs):
                text = para.text
                if "{{LOOP:" in text:
                    start = text.find("{{LOOP:") + 7
                    end = text.find("}}", start)
                    loop_var = text[start:end].strip()
                    loop_marker_idx = i
                elif "{{titulo}}" in text and loop_marker_idx is not None:
                    titulo_idx = i
                elif "{{paragrafo}}" in text and loop_marker_idx is not None:
                    paragrafo_idx = i
                    # Não damos break aqui, continuamos para ver se achamos mais coisas, 
                    # mas o loop funciona mesmo sem paragrafo_idx
            
            # Validação: Precisa pelo menos do Loop e do Título
            if loop_marker_idx is not None and titulo_idx is not None and loop_var:
                if loop_var in context_data and isinstance(context_data[loop_var], list):
                    loop_data = context_data[loop_var]
                    
                    parent = doc.paragraphs[loop_marker_idx]._element.getparent()
                    loop_marker_elem = doc.paragraphs[loop_marker_idx]._element
                    titulo_elem = doc.paragraphs[titulo_idx]._element
                    
                    # Parágrafo agora é opcional
                    paragrafo_elem = None
                    if paragrafo_idx is not None:
                        paragrafo_elem = doc.paragraphs[paragrafo_idx]._element
                    
                    insert_pos = parent.index(loop_marker_elem)
                    
                    # Remove templates originais
                    parent.remove(loop_marker_elem)
                    parent.remove(titulo_elem)
                    if paragrafo_elem is not None:
                        parent.remove(paragrafo_elem)
                    
                    # Expansão do Loop
                    for item in loop_data:
                        # Clona Título
                        new_titulo = deepcopy(titulo_elem)
                        parent.insert(insert_pos, new_titulo)
                        insert_pos += 1
                        
                        # Clona Parágrafo (SÓ SE EXISTIR)
                        if paragrafo_elem is not None:
                            new_para = deepcopy(paragrafo_elem)
                            parent.insert(insert_pos, new_para)
                            insert_pos += 1
                    
                    # Substituição de Texto nos novos itens
                    t_count = 0
                    p_count = 0
                    
                    # Precisamos iterar novamente pelos parágrafos do documento atualizado
                    for para in doc.paragraphs:
                        if "{{titulo}}" in para.text and t_count < len(loop_data):
                            val = loop_data[t_count].get("titulo", "")
                            _replace_in_runs(para, "{{titulo}}", str(val))
                            t_count += 1
                        elif "{{paragrafo}}" in para.text and p_count < len(loop_data):
                            val = loop_data[p_count].get("paragrafo", "")
                            _replace_in_runs(para, "{{paragrafo}}", str(val))
                            p_count += 1
            
            # STEP 2: Replace simple variables in body
            for para in doc.paragraphs:
                for key, value in context_data.items():
                    if not isinstance(value, (list, dict)):
                        placeholder = f"{{{{{key}}}}}"
                        _replace_in_runs(para, placeholder, str(value))
                        
            try:
                logging.info("Chamando _process_dynamic_table_placeholder agora...") # <--- ADICIONE ISSO
                _process_dynamic_table_placeholder(doc, context_data, "tabela_dinamica")
            except Exception as e:
                logging.error(f"Erro na chamada da tabela: {e}")
            
            # STEP 3: Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for key, value in context_data.items():
                                if not isinstance(value, (list, dict)):
                                    placeholder = f"{{{{{key}}}}}"
                                    _replace_in_runs(para, placeholder, str(value))
            
            # STEP 4: Replace in headers
            for section in doc.sections:
                header = section.header
                for para in header.paragraphs:
                    for key, value in context_data.items():
                        if not isinstance(value, (list, dict)):
                            placeholder = f"{{{{{key}}}}}"
                            _replace_in_runs(para, placeholder, str(value))
                
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for key, value in context_data.items():
                                    if not isinstance(value, (list, dict)):
                                        placeholder = f"{{{{{key}}}}}"
                                        _replace_in_runs(para, placeholder, str(value))
            

            for section in doc.sections:
                # Lista com os 3 tipos de rodapé (Padrão, Primeira Pág, Pares)
                footers = [
                    section.footer, 
                    section.first_page_footer, 
                    section.even_page_footer
                ]
                
                for footer in footers:
                    if not footer: continue
                    
                    # AQUI ESTÁ A CORREÇÃO:
                    # Varre todo o XML do rodapé procurando por parágrafos (<w:p>)
                    # Isso encontra parágrafos escondidos dentro de tabelas e CAIXAS DE TEXTO
                    for element in footer._element.iter():
                        if isinstance(element, CT_P):
                            # Transforma o XML achado em um objeto Parágrafo do Python
                            para = Paragraph(element, footer)
                            
                            # Agora usamos a função inteligente que conserta "texto quebrado"
                            for key, value in context_data.items():
                                if not isinstance(value, (list, dict)):
                                    placeholder = f"{{{{{key}}}}}"
                                    
                                    # Verifica se o placeholder está no texto (mesmo que quebrado)
                                    if placeholder in para.text:
                                        _replace_in_runs(para, placeholder, str(value))
                
                # Process XML for textboxes
                try:
                    from lxml import etree
                    
                    footer_xml = etree.tostring(footer._element, encoding='unicode')
                    
                    if "{{" in footer_xml and "}}" in footer_xml:
                        modified = False
                        for key, value in context_data.items():
                            if not isinstance(value, (list, dict)):
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in footer_xml:
                                    footer_xml = footer_xml.replace(placeholder, str(value))
                                    modified = True
                        
                        if modified:
                            new_footer_element = etree.fromstring(footer_xml.encode('utf-8'))
                            parent = footer._element.getparent()
                            if parent is not None:
                                index = list(parent).index(footer._element)
                                parent.remove(footer._element)
                                parent.insert(index, new_footer_element)
                except:
                    pass
            
            settings = doc.settings.element
            update_fields = OxmlElement('w:updateFields')
            update_fields.set(qn('w:val'), 'true')
            settings.append(update_fields)
            
            doc.save(output_path)
            return True, f"Template filled successfully. Output saved to {output_path}"
        except Exception as e:
            import traceback
            return False, f"Error filling template: {str(e)}\n{traceback.format_exc()}"
    
    # Execute the blocking operation in a thread pool to avoid blocking the async loop
    try:
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            success, message = await loop.run_in_executor(executor, _fill_template)
            return message
    except Exception as e:
        return f"Failed to fill document template: {str(e)}"
