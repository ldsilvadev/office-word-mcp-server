import asyncio
import os
from docx import Document
from word_document_server.tools.content_tools import add_section_with_inherited_formatting

async def reproduce():
    filename = "test_spacing_trailing.docx"
    if os.path.exists(filename):
        os.remove(filename)
    
    # Create initial document with a trailing empty paragraph
    doc = Document()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Some content here.")
    doc.add_paragraph("") # Trailing empty paragraph
    doc.save(filename)
    
    print(f"Created {filename} with initial content and trailing empty paragraph.")
    
    # Add a new section
    print("Adding new section...")
    result = await add_section_with_inherited_formatting(filename, "Section 2", "Content for section 2")
    print(result)
    
    # Inspect the document structure
    doc = Document(filename)
    print("\nDocument paragraphs:")
    for i, p in enumerate(doc.paragraphs):
        print(f"{i}: '{p.text}' (Style: {p.style.name})")

if __name__ == "__main__":
    asyncio.run(reproduce())
