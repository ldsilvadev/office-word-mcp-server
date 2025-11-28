import asyncio
import os
import sys
sys.path.append(os.getcwd())
from docx import Document
from word_document_server.tools.content_tools import add_section_with_inherited_formatting, add_heading

async def test_inheritance():
    filename = "test_inheritance.docx"
    
    # Create a clean document
    if os.path.exists(filename):
        os.remove(filename)
    
    doc = Document()
    # Add a heading with a specific style (Heading 1 is standard)
    doc.add_heading("Section 1", level=1)
    
    # Add a paragraph with a specific style (e.g., 'Quote' which usually exists)
    p = doc.add_paragraph("Content 1")
    p.style = "Quote"
    
    doc.save(filename)
    
    print(f"Created {filename} with 'Section 1' and paragraph style 'Quote'")
    
    # Test adding a new section
    print("Adding 'Section 2'...")
    result = await add_section_with_inherited_formatting(
        filename=filename, 
        title="Section 2", 
        paragraph_text="This is the content for section 2.",
        table_data=[["Header A", "Header B"], ["Cell 1", "Cell 2"]]
    )
    print(result)
    
    # Verify
    doc = Document(filename)
    
    # Check headings
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    print(f"Found {len(headings)} headings:")
    for h in headings:
        print(f" - {h.text} ({h.style.name})")
        
    # Check body paragraphs
    # We expect: Heading 1, Quote, Heading 1, Quote (inherited)
    # Note: doc.paragraphs includes all paragraphs.
    
    paras = doc.paragraphs
    if len(paras) >= 4:
        p2 = paras[1] # Content 1
        p4 = paras[3] # Content 2
        print(f"Paragraph 1 style: {p2.style.name}")
        print(f"Paragraph 2 style: {p4.style.name}")
        
        if p4.style.name == "Quote":
             print("SUCCESS: Body style 'Quote' inherited correctly.")
        else:
             print(f"FAILURE: Body style not inherited. Expected 'Quote', got '{p4.style.name}'")

    if len(headings) == 2 and headings[1].text == "Section 2" and headings[1].style.name == headings[0].style.name:
        print("SUCCESS: Heading style inherited correctly.")
    else:
        print("FAILURE: Heading style not inherited or heading missing.")

if __name__ == "__main__":
    asyncio.run(test_inheritance())
