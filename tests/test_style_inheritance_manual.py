import asyncio
import os
import sys
from docx import Document

# Add project root to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from word_document_server.tools.content_tools import add_paragraph
from word_document_server.tools.document_tools import create_document

async def test_add_paragraph_style():
    filename = "test_style_inheritance.docx"
    
    # 1. Create document
    print(f"Creating {filename}...")
    await create_document(filename)
    
    # 2. Add a paragraph with a specific style (e.g., 'Quote' or 'List Paragraph')
    # We need to ensure the style exists or use a standard one. 'Quote' is standard.
    print("Adding initial paragraph with 'Quote' style...")
    await add_paragraph(filename, "Initial Quote", style="Quote")
    
    # 3. Add a new paragraph WITHOUT specifying style
    print("Adding second paragraph (should inherit 'Quote')...")
    await add_paragraph(filename, "Inherited Quote")
    
    # 4. Verify styles
    doc = Document(filename)
    p1 = doc.paragraphs[0]
    p2 = doc.paragraphs[1]
    
    print(f"P1 Style: {p1.style.name}")
    print(f"P2 Style: {p2.style.name}")
    
    if p2.style.name == "Quote":
        print("SUCCESS: Style inherited correctly.")
    else:
        print(f"FAILURE: Expected 'Quote', got '{p2.style.name}'")

    # 5. Test Heading exception
    # Add a Heading
    print("Adding Heading...")
    doc.add_heading("My Heading", level=1)
    doc.save(filename)
    
    # Add paragraph after heading (should NOT inherit Heading 1)
    print("Adding paragraph after Heading (should NOT inherit Heading)...")
    await add_paragraph(filename, "Normal Text")
    
    doc = Document(filename)
    p_last = doc.paragraphs[-1]
    print(f"Last Para Style: {p_last.style.name}")
    
    if p_last.style.name != "Heading 1":
         print("SUCCESS: Heading style NOT inherited.")
    else:
         print("FAILURE: Heading style WAS inherited.")

    # Cleanup
    if os.path.exists(filename):
        os.remove(filename)

if __name__ == "__main__":
    asyncio.run(test_add_paragraph_style())
